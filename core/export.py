# core/export.py

import html, io, difflib
from docx import Document
from docx.shared import RGBColor

# ---------- Inline Word Diff (HTML) ----------
def inline_word_diff_html(a, b):
    a_words = a.split()
    b_words = b.split()
    matcher = difflib.SequenceMatcher(a=a_words, b=b_words)
    out = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            out.append(html.escape(" ".join(a_words[i1:i2])))
        elif tag == "delete":
            deleted = html.escape(" ".join(a_words[i1:i2]))
            out.append(f'<del style="color:#a00;text-decoration:line-through;">{deleted}</del>')
        elif tag == "insert":
            inserted = html.escape(" ".join(b_words[j1:j2]))
            out.append(f'<ins style="color:green;text-decoration:none;">{inserted}</ins>')
        elif tag == "replace":
            deleted = html.escape(" ".join(a_words[i1:i2]))
            inserted = html.escape(" ".join(b_words[j1:j2]))
            out.append(f'<del style="color:#a00;text-decoration:line-through;">{deleted}</del>')
            out.append(f'<ins style="color:green;text-decoration:none;">{inserted}</ins>')
    return " ".join(out)

# ---------- Export Tracked Changes (HTML) ----------
def export_tracked_html(export_items):
    parts = ["<html><meta charset='utf-8'><body><h1>Negotiation Draft</h1>"]
    for item in export_items:
        parts.append(f"<h3>Clause {item['id']} — Persona: {html.escape(item['persona'])} — Style: {html.escape(item['style'])}</h3>")
        parts.append(f"<p><strong>Original:</strong><br><pre>{html.escape(item['original'])}</pre></p>")
        counter = item.get("counter", "")
        if counter:
            parts.append("<p><strong>Proposed:</strong><br>")
            parts.append(inline_word_diff_html(item['original'], counter))
            parts.append("</p>")
        if item.get("simulated_reply"):
            parts.append(f"<p><em>Simulated reply:</em><br><pre>{html.escape(item['simulated_reply'])}</pre></p>")
        parts.append("<hr/>")
    parts.append("</body></html>")
    return "\n".join(parts)

# ---------- Export Tracked Changes (DOCX) ----------
def build_docx_with_diffs(export_items, max_original_chars=4000):
    doc = Document()
    doc.core_properties.title = "Negotiation Draft"
    doc.core_properties.comments = f"Generated: {datetime.datetime.utcnow().isoformat()}Z"
    for item in export_items:
        doc.add_paragraph(f"Clause {item['id']} — Persona: {item['persona']} | Style: {item['style']}")
        p_orig = doc.add_paragraph()
        p_orig.add_run("Original:\n").bold = True
        orig_text = item["original"][:max_original_chars]
        p_orig.add_run(orig_text)
        counter_text = item.get("counter", "") or ""
        if counter_text:
            p_counter_label = doc.add_paragraph()
            p_counter_label.add_run("\nCounter-Proposal:\n").bold = True
            a_words = orig_text.split()
            b_words = counter_text.split()
            matcher = difflib.SequenceMatcher(a=a_words, b=b_words)
            p_counter = doc.add_paragraph()
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == "equal":
                    p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                elif tag == "insert":
                    run = p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                elif tag == "delete":
                    run = p_counter.add_run(" " + " ".join(a_words[i1:i2]))
                    run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
                    run.font.strike = True
                elif tag == "replace":
                    run_del = p_counter.add_run(" " + " ".join(a_words[i1:i2]))
                    run_del.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
                    run_del.font.strike = True
                    run_ins = p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                    run_ins.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        if item.get("simulated_reply"):
            p_sim = doc.add_paragraph()
            p_sim.add_run("\nSimulated Counterparty Reply:\n").bold = True
            r = p_sim.add_run(item["simulated_reply"])
            r.italic = True
        doc.add_paragraph("---")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
