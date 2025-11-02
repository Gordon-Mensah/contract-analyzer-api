# Home.py — Contract Intelligence Platform (Simplified & Roadmap-Aligned)

# ---------- Imports ----------
import streamlit as st
import os, json, datetime, html, re, io, tempfile, time
import numpy as np
import difflib
from collections import Counter
from hashlib import sha256
from docx import Document
from docx.shared import RGBColor
import pdfplumber
import matplotlib.pyplot as plt
from sentence_transformers import SentenceTransformer
from transformers import pipeline
from deep_translator import GoogleTranslator
from langchain.text_splitter import RecursiveCharacterTextSplitter
import chromadb
from chromadb.config import Settings
import diskcache as dc
import joblib
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import roc_auc_score, precision_recall_fscore_support
from sklearn.model_selection import train_test_split

# ---------- Helpers: cache keys ----------
def mkhash(*args) -> str:
    h = sha256()
    for a in args:
        if isinstance(a, str):
            h.update(a.encode("utf-8"))
        else:
            h.update(json.dumps(a, sort_keys=True, default=str).encode("utf-8"))
    return h.hexdigest()

# ---------- App Config ----------
st.set_page_config(page_title="Contract Intelligence", page_icon="📄", layout="wide")

# ---------- Constants ----------
CACHE_DIR = ".cache_disk"
CACHE_TTL = 60 * 60 * 24
STATE_FILE = ".session_state.json"
PERSONA_FILE = "personas.json"
FEEDBACK_FILE = "feedback_store.json"
WEIGHTS_FILE = "ranking_weights.json"
MODEL_FILE = "feedback_model.joblib"
DEFAULT_WEIGHTS = {"w_risk": 2.0, "w_sim": 1.0, "w_len": 0.5}
AUTO_TRAIN_THRESHOLD = 50
MIN_SAMPLES_FOR_MODEL = 20
MIN_POSITIVE_FOR_MODEL = 5

# ---------- Disk Cache ----------
cache = dc.Cache(CACHE_DIR)

# ---------- Cached Models ----------
@st.cache_resource
def get_summarizer():
    return pipeline("summarization", model="t5-small")

@st.cache_resource
def get_sentence_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def get_embedder():
    return get_sentence_model()

# ---------- ChromaDB Setup ----------
client = chromadb.Client(Settings())
def get_collection_for_text(text: str):
    h = str(abs(hash(text)))[:8]
    return client.get_or_create_collection(name=f"contract_chunks_{h}")

def clear_collection(collection):
    try:
        client.delete_collection(name=collection.name)
        return client.get_or_create_collection(name=collection.name)
    except:
        return collection

# ---------- Session Initialization ----------
def load_personas():
    if os.path.exists(PERSONA_FILE):
        with open(PERSONA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def load_session_state():
    if os.path.exists(STATE_FILE):
        with open(STATE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        st.session_state.negotiation_text = data.get("contract_text", "")
        st.session_state.labeled_chunks = data.get("labeled_chunks", [])
        st.session_state.neg_turns = data.get("neg_turns", [])
        st.session_state.neg_personas = data.get("neg_personas", {})
        st.session_state.neg_counters = data.get("neg_counters", {})

# Initialize session state
_defaults = {
    "neg_personas": {},
    "neg_counters": {},
    "neg_simulated": {},
    "neg_turns": [],
    "labeled_chunks": [],
    "negotiation_text": "",
    "saved_personas": load_personas(),
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v
load_session_state()
# ---------- Contract Upload ----------
def load_contract(file):
    if file.name.lower().endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif file.name.lower().endswith(".docx"):
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return file.read().decode("utf-8", errors="ignore")

def chunk_contract(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_text(text)

# ---------- Translation ----------
def translate_to_hungarian(text):
    try:
        return GoogleTranslator(source='auto', target='hu').translate(text)
    except Exception:
        return ""

        # ---------- Inline Risk Highlighting ----------
RISK_HIGHLIGHTS = {
    "indemnify": "background:#ffd6d6",
    "exclusive": "background:#ffd6d6",
    "penalty": "background:#ffd6d6",
    "binding": "background:#fff0b3",
    "termination": "background:#fff0b3"
}

def highlight_risks(text):
    safe = html.escape(text)
    for term, style in RISK_HIGHLIGHTS.items():
        safe = re.sub(fr"(?i)\b({re.escape(term)})\b", rf"<span style='{style};padding:2px;border-radius:3px'>\1</span>", safe)
    return safe

# ---------- Risk Scoring ----------
def compute_risk_score(text: str) -> int:
    _, risk_level = label_clause(text)
    mapping = {"Low": 1, "Medium": 2, "High": 3}
    return mapping.get(risk_level, 1)


# ---------- Clause Labeling ----------
def label_clause(text):
    clause_type = "Other"
    risk_level = "Low"
    keywords = {
        "Termination": ["terminate", "termination", "cancel", "end of agreement"],
        "Confidentiality": ["confidential", "non-disclosure", "privacy", "secret"],
        "Non-compete": ["non-compete", "competition", "exclusive", "restrict"],
        "Indemnity": ["indemnify", "liability", "hold harmless", "damages"],
        "Payment": ["payment", "fee", "compensation", "invoice", "cost"],
        "Jurisdiction": ["governing law", "jurisdiction", "venue", "court"],
        "Timeline": ["effective", "duration", "term", "deadline", "return within", "test period", "expires", "termination date"]
    }
    risks = {
        "High": ["penalty", "irreversible", "binding", "exclusive", "indemnify"],
        "Medium": ["termination", "non-compete", "governing law", "confidential"],
        "Low": ["notice", "duration", "payment", "invoice"]
    }
    text_l = text.lower()
    for label, terms in keywords.items():
        if any(term in text_l for term in terms):
            clause_type = label
            break
    for level, terms in risks.items():
        if any(term in text_l for term in terms):
            risk_level = level
            break
    return clause_type, risk_level

def is_red_flag(text, risk_level):
    red_terms = ["indemnify", "exclusive", "binding", "penalty", "irreversible"]
    return risk_level == "High" and any(term in text.lower() for term in red_terms)

def format_badges(clause_type, risk_level):
    risk_colors = {"High": "🔴 High Risk", "Medium": "🟠 Medium Risk", "Low": "🟢 Low Risk"}
    return f"📌 **{clause_type}** | {risk_colors.get(risk_level, '⚪ Unknown Risk')}"

# ---------- Rewrite Suggestions ----------
def summarize_clause(text):
    key = "summarize:" + mkhash(text)
    res = cache.get(key)
    if res is not None:
        return res
    summarizer = get_summarizer()
    try:
        out = summarizer(text, max_length=60, min_length=20, do_sample=False)
        s = out[0]["summary_text"] if isinstance(out, list) else out.get("summary_text", "")
    except Exception:
        s = text[:200] + "..." if len(text) > 200 else text
    cache.set(key, s, expire=CACHE_TTL)
    return s

def generate_rewrite_candidates(original_text, persona, style, n_candidates=3):
    candidates = []
    base_prompt = f"Rewrite the following contract clause for negotiation. Persona: {persona}. Style: {style}.\n\nClause:\n{original_text}"
    for i in range(n_candidates):
        prompt = base_prompt + f"\n\nCandidate variation: {i+1}"
        key = "rewrite_candidate:" + mkhash(prompt)
        cached = cache.get(key)
        if cached:
            candidates.append(cached)
            continue
        try:
            out = get_summarizer()(prompt, max_length=120, min_length=30, do_sample=True, top_k=50, top_p=0.95)
            text = out[0]["summary_text"].strip() if isinstance(out, list) else out.get("summary_text", "").strip()
        except Exception:
            text = original_text[:200] + "..."
        cache.set(key, text, expire=CACHE_TTL)
        candidates.append(text)
    return candidates

# ---------- Feedback Logging ----------
def log_feedback(clause_index, original, candidate, accepted, meta, score, action):
    store = load_feedback_store()
    entry = {
        "timestamp": datetime.datetime.utcnow().isoformat() + "Z",
        "clause_index": clause_index,
        "accepted": bool(accepted),
        "action": action,
        "score": float(score),
        "meta": meta,
        "original_len": len(original),
        "candidate_len": len(candidate),
        "candidate_text": candidate[:1000]
    }
    store.append(entry)
    save_feedback_store(store)
    if len(store) >= AUTO_TRAIN_THRESHOLD:
        try:
            _ = train_feedback_model(min_samples=MIN_SAMPLES_FOR_MODEL)
        except Exception:
            pass

def load_feedback_store():
    if os.path.exists(FEEDBACK_FILE):
        with open(FEEDBACK_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_feedback_store(store):
    with open(FEEDBACK_FILE, "w", encoding="utf-8") as f:
        json.dump(store, f, ensure_ascii=False, indent=2)

# ---------- Ranking ----------
def embed_text(text):
    key = "embed_text:" + mkhash(text)
    res = cache.get(key)
    if res is not None:
        return np.array(res, dtype=float)
    model = get_embedder()
    vec = model.encode([text])[0]
    cache.set(key, vec, expire=CACHE_TTL)
    return np.array(vec, dtype=float)

def cosine_similarity(a, b):
    if np.linalg.norm(a) == 0 or np.linalg.norm(b) == 0:
        return 0.0
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))

def score_candidate_heuristic(original, candidate):
    orig_risk = compute_risk_score(original)
    cand_risk = compute_risk_score(candidate)
    risk_delta = orig_risk - cand_risk
    sim = cosine_similarity(embed_text(original), embed_text(candidate))
    len_orig = len(original)
    len_cand = len(candidate)
    length_penalty = max(0, 1 - abs(len_cand - len_orig) / max(10, len_orig))
    w = load_weights()
    score = (w["w_risk"] * risk_delta) + (w["w_sim"] * sim) + (w["w_len"] * length_penalty)
    return float(score), {"risk_delta": risk_delta, "similarity": sim, "len_penalty": length_penalty}

def load_weights():
    if os.path.exists(WEIGHTS_FILE):
        with open(WEIGHTS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return DEFAULT_WEIGHTS.copy()

def save_weights(w):
    with open(WEIGHTS_FILE, "w", encoding="utf-8") as f:
        json.dump(w, f, ensure_ascii=False, indent=2)

# ---------- Classifier ----------
def build_training_set():
    store = load_feedback_store()
    X, y = [], []
    for e in store:
        meta = e.get("meta", {})
        feat = [
            meta.get("risk_delta", 0.0),
            meta.get("similarity", 0.0),
            meta.get("len_penalty", 0.0),
            e.get("original_len", 0),
            e.get("candidate_len", 0)
        ]
        X.append(feat)
        y.append(1 if e.get("accepted") else 0)
    return np.array(X), np.array(y)

def train_feedback_model(min_samples=MIN_SAMPLES_FOR_MODEL):
    X, y = build_training_set()
    if len(X) < min_samples or sum(y) < MIN_POSITIVE_FOR_MODEL:
        return {"ok": False, "reason": "not enough data"}
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, stratify=y)
    clf = LogisticRegression(max_iter=1000)
    clf.fit(X_train, y_train)
    joblib.dump(clf, MODEL_FILE)
    return {"ok": True}

def load_feedback_model():
    if os.path.exists(MODEL_FILE):
        return joblib.load(MODEL_FILE)
    return None

def predict_accept_prob(meta):
    clf = load_feedback_model()
    if clf is None:
        return None
    feat = np.array([
        meta.get("risk_delta", 0.0),
        meta.get("similarity", 0.0),
        meta.get("len_penalty", 0.0),
        meta.get("original_len", 0),
        meta.get("candidate_len", 0)
    ]).reshape(1, -1)
    try:
        return float(clf.predict_proba(feat)[0, 1])
    except:
        return None
# ---------- Negotiation Simulation ----------
def add_turn(author, text, persona, role):
    turn_id = len(st.session_state.neg_turns) + 1
    emb = embed_text(text)
    st.session_state.neg_turns.append({
        "turn": turn_id,
        "author": author,
        "text": text,
        "persona": persona,
        "role": role,
        "embedding": emb.tolist(),
        "timestamp": datetime.datetime.utcnow().isoformat() + "Z"
    })

def auto_negotiate_simulation(clause_text, persona, turns=3, stop_threshold=0.9, pause=0.5):
    history = []
    current = clause_text
    for t in range(turns):
        prompt = f"Rewrite clause for persona {persona}:\n\n{current}"
        candidate = summarize_clause(prompt)
        history.append(("you", candidate))
        add_turn("you", candidate, persona, "offer")
        if t > 0:
            a = embed_text(history[-2][1])
            b = embed_text(history[-1][1])
            if cosine_similarity(a, b) >= stop_threshold:
                break
        time.sleep(pause)
        sim_prompt = f"Respond as counterparty to: {candidate}"
        sim = summarize_clause(sim_prompt)
        history.append(("counterparty", sim))
        add_turn("counterparty", sim, persona, "reply")
        time.sleep(pause)
    return history

# ---------- Tracked Changes (HTML) ----------
def inline_word_diff_html(a, b):
    a_words = a.split()
    b_words = b.split()
    matcher = difflib.SequenceMatcher(a=a_words, b=b_words)
    out = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            out.append(html.escape(" ".join(a_words[i1:i2])))
        elif tag == "delete":
            deleted = html.escape(" ".join(a_words[i1:i2]))
            out.append(f'<del style="color:#a00;text-decoration:line-through;">{deleted}</del>')
        elif tag == "insert":
            inserted = html.escape(" ".join(b_words[j1:j2]))
            out.append(f'<ins style="color:green;text-decoration:none;">{inserted}</ins>')
        elif tag == "replace":
            deleted = html.escape(" ".join(a_words[i1:i2]))
            inserted = html.escape(" ".join(b_words[j1:j2]))
            out.append(f'<del style="color:#a00;text-decoration:line-through;">{deleted}</del>')
            out.append(f'<ins style="color:green;text-decoration:none;">{inserted}</ins>')
    return " ".join(out)

def export_tracked_html(export_items):
    parts = ["<html><meta charset='utf-8'><body><h1>Negotiation Draft</h1>"]
    for item in export_items:
        parts.append(f"<h3>Clause {item['id']} — Persona: {html.escape(item['persona'])} — Style: {html.escape(item['style'])}</h3>")
        parts.append(f"<p><strong>Original:</strong><br><pre>{html.escape(item['original'])}</pre></p>")
        counter = item.get("counter", "")
        if counter:
            parts.append("<p><strong>Proposed:</strong><br>")
            parts.append(inline_word_diff_html(item['original'], counter))
            parts.append("</p>")
        if item.get("simulated_reply"):
            parts.append(f"<p><em>Simulated reply:</em><br><pre>{html.escape(item['simulated_reply'])}</pre></p>")
        parts.append("<hr/>")
    parts.append("</body></html>")
    return "\n".join(parts)

# ---------- Tracked Changes (DOCX) ----------
def build_docx_with_diffs(export_items, max_original_chars=4000):
    doc = Document()
    doc.core_properties.title = "Negotiation Draft"
    doc.core_properties.comments = f"Generated: {datetime.datetime.utcnow().isoformat()}Z"
    for item in export_items:
        doc.add_paragraph(f"Clause {item['id']} — Persona: {item['persona']} | Style: {item['style']}")
        p_orig = doc.add_paragraph()
        p_orig.add_run("Original:\n").bold = True
        orig_text = item["original"][:max_original_chars]
        p_orig.add_run(orig_text)
        counter_text = item.get("counter", "") or ""
        if counter_text:
            p_counter_label = doc.add_paragraph()
            p_counter_label.add_run("\nCounter-Proposal:\n").bold = True
            a_words = orig_text.split()
            b_words = counter_text.split()
            matcher = difflib.SequenceMatcher(a=a_words, b=b_words)
            p_counter = doc.add_paragraph()
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == "equal":
                    p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                elif tag == "insert":
                    run = p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                elif tag == "delete":
                    run = p_counter.add_run(" " + " ".join(a_words[i1:i2]))
                    run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
                    run.font.strike = True
                elif tag == "replace":
                    run_del = p_counter.add_run(" " + " ".join(a_words[i1:i2]))
                    run_del.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
                    run_del.font.strike = True
                    run_ins = p_counter.add_run(" " + " ".join(b_words[j1:j2]))
                    run_ins.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        if item.get("simulated_reply"):
            p_sim = doc.add_paragraph()
            p_sim.add_run("\nSimulated Counterparty Reply:\n").bold = True
            r = p_sim.add_run(item["simulated_reply"])
            r.italic = True
        doc.add_paragraph("---")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------- Sidebar ----------
st.sidebar.title("📄 Contract Intelligence")
st.sidebar.markdown("Build smarter contracts with zero budget")

with st.sidebar.expander("📁 Upload Contract"):
    uploaded_file = st.file_uploader("Choose a contract file", type=["pdf", "docx", "txt"])
    if uploaded_file:
        text = load_contract(uploaded_file)
        st.session_state.negotiation_text = text
        st.success("Contract loaded.")

    with st.sidebar.expander("🌐 Import from Link"):
        contract_url = st.text_input("Paste contract URL (PDF, DOCX, or TXT)")
    if contract_url:
        try:
            response = requests.get(contract_url)
            if response.status_code == 200:
                content_type = response.headers.get("Content-Type", "")
                temp_file = tempfile.NamedTemporaryFile(delete=False)
                temp_file.write(response.content)
                temp_file.flush()
                temp_file.seek(0)
                if "pdf" in content_type:
                    text = load_contract(temp_file)
                elif "word" in content_type or contract_url.lower().endswith(".docx"):
                    text = load_contract(temp_file)
                else:
                    text = response.text
                st.session_state.negotiation_text = text
                st.success("Contract imported from link.")
            else:
                st.warning(f"Failed to fetch document. Status code: {response.status_code}")
        except Exception as e:
            st.error(f"Error fetching document: {e}")
summarize_enabled = st.sidebar.checkbox("Summarize clauses", value=True)
translate_enabled = st.sidebar.checkbox("Translate to Hungarian", value=False)


with st.sidebar.expander("🧠 Persona Settings"):
    persona = st.text_input("Persona name", value="Startup Founder")
    style = st.selectbox("Rewrite style", ["Plain English", "Legalese", "Assertive", "Concise", "Friendly"])
    st.session_state.neg_personas["default"] = {"persona": persona, "style": style}

with st.sidebar.expander("📊 Advanced Controls"):
    ranking_mode = st.selectbox("Ranking mode", ["Auto (Classifier)", "Heuristic Only"])
    if st.button("Improve Suggestions (Train Model)"):
        train_feedback_model()
        st.toast("Model training complete.")

if st.sidebar.button("💾 Save Session"):
    save_session_state()

if st.sidebar.button("📂 Load Session"):
    load_session_state()

# ---------- Candidate Presentation UI ----------
def present_top_candidates_ui(original_text, clause_index, persona, style, ranking_mode):
    st.markdown("### ✨ Suggested Counter-Proposals")
    candidates = generate_rewrite_candidates(original_text, persona, style, n_candidates=3)
    scored = []
    for c in candidates:
        heur_score, meta = score_candidate_heuristic(original_text, c)
        meta_ext = meta.copy()
        meta_ext["original_len"] = len(original_text)
        meta_ext["candidate_len"] = len(c)
        prob = predict_accept_prob(meta_ext) if ranking_mode.startswith("Auto") else None
        rank_score = prob if prob is not None else heur_score
        scored.append({
            "text": c,
            "heur_score": heur_score,
            "meta": meta_ext,
            "prob": prob,
            "rank_score": rank_score
        })
    scored = sorted(scored, key=lambda x: x["rank_score"], reverse=True)

    for rank, item in enumerate(scored, start=1):
        display_score = item["prob"] if item["prob"] is not None else item["heur_score"]
        st.markdown(f"#### #{rank} — Score: **{display_score:.3f}** {'(prob)' if item['prob'] is not None else '(heuristic)'}")
        st.markdown(f"- **Similarity:** {item['meta']['similarity']:.3f}; **Risk Delta:** {item['meta']['risk_delta']:.1f}; **Length Score:** {item['meta']['len_penalty']:.3f}")
        diff_html = inline_word_diff_html(original_text, item["text"])
        st.markdown(diff_html, unsafe_allow_html=True)

        col_a, col_b = st.columns([3, 1])
        with col_a:
            new_text = st.text_area(f"Edit Candidate {clause_index}_{rank}", value=item["text"], key=f"candidate_edit_{clause_index}_{rank}", height=120)
        with col_b:
            if st.button(f"✅ Accept #{rank}", key=f"accept_{clause_index}_{rank}"):
                if "labeled_chunks" in st.session_state and 0 <= clause_index < len(st.session_state.labeled_chunks):
                    st.session_state.labeled_chunks[clause_index]["text"] = new_text
                    st.session_state.neg_counters[f"counter_{clause_index}"] = new_text
                    add_turn("you", new_text, persona, "accepted_offer")
                    log_feedback(clause_index, original_text, new_text, True, item["meta"], display_score, "replace")
                    st.success(f"Accepted candidate #{rank} for Clause {clause_index + 1}")


# ---------- Main UI ----------
st.title("📄 Contract Intelligence Platform")
st.markdown("Upload a contract, analyze clauses, review rewrite suggestions, and simulate negotiation.")

if st.session_state.negotiation_text:
    st.subheader("📜 Original Contract Text")
    st.text_area("Contract", value=st.session_state.negotiation_text, height=200)

    if st.button("🔍 Analyze Clauses"):
        chunks = chunk_contract(st.session_state.negotiation_text)
        labeled = []
        for i, chunk in enumerate(chunks):
            clause_type, risk_level = label_clause(chunk)
            summary = summarize_clause(chunk) if summarize_enabled else ""
            translated = translate_to_hungarian(chunk) if translate_enabled else ""
            labeled.append({
                    "id": i,
                    "text": chunk,
                    "type": clause_type,
                    "risk": risk_level,
                    "summary": summary,
                    "translated": translated
    })
        st.session_state.labeled_chunks = labeled
        st.success(f"{len(labeled)} clauses analyzed.")

if st.session_state.labeled_chunks:
    st.subheader("🧩 Clause Review")
    for i, clause in enumerate(st.session_state.labeled_chunks):
        with st.expander(f"Clause {i+1}: {format_badges(clause['type'], clause['risk'])}"):
            st.markdown(highlight_risks(clause["text"]), unsafe_allow_html=True)
            st.markdown(f"**Summary:** {clause['summary']}")
            st.markdown(f"**Hungarian:** {clause['translated']}")
            present_top_candidates_ui(clause["text"], i, persona, style, ranking_mode)

    if st.button("📤 Export Tracked Changes (HTML)"):
        html_data = export_tracked_html([
            {
                "id": c["id"],
                "original": c["text"],
                "counter": st.session_state.neg_counters.get(f"counter_{c['id']}", ""),
                "persona": persona,
                "style": style,
                "simulated_reply": st.session_state.neg_simulated.get(f"reply_{c['id']}", "")
            } for c in st.session_state.labeled_chunks
        ])
        st.download_button("Download HTML", data=html_data, file_name="tracked_changes.html")

    if st.button("📄 Export Tracked Changes (DOCX)"):
        docx_data = build_docx_with_diffs([
            {
                "id": c["id"],
                "original": c["text"],
                "counter": st.session_state.neg_counters.get(f"counter_{c['id']}", ""),
                "persona": persona,
                "style": style,
                "simulated_reply": st.session_state.neg_simulated.get(f"reply_{c['id']}", "")
            } for c in st.session_state.labeled_chunks
        ])
        st.download_button("Download DOCX", data=docx_data, file_name="tracked_changes.docx")
